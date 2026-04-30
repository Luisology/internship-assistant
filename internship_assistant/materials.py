"""Application material templates — context-aware, DOCX-capable.

V11 improvements:
- Uses company_context for personalization
- Extracts relevant resume lines (never dumps full resume)
- DOCX as primary format, TXT as fallback
- Honest resume suggestions with clear disclaimers
- Follow-up and thank-you message templates
"""

import os
import re
from datetime import date


# ── Filename helpers ──────────────────────────────────────────────────────────

def safe_filename(text):
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return cleaned or "untitled"


def get_applicant_name(resume_path):
    if not os.path.exists(resume_path):
        return "[Your Name]"
    with open(resume_path) as f:
        for line in f:
            line = line.strip()
            if line:
                return line
    return "[Your Name]"


# ── Resume helpers ────────────────────────────────────────────────────────────

def extract_relevant_resume_lines(resume_text, skills, max_lines=3):
    """Return bullet-style resume lines that mention any matched skill.

    Never returns more than max_lines; never returns the full resume.
    """
    if not resume_text or not skills:
        return []
    skills_lower = [s.lower() for s in skills]
    relevant = []
    for raw_line in resume_text.splitlines():
        line = raw_line.strip()
        if len(line) < 25:
            continue
        if any(sk in line.lower() for sk in skills_lower):
            clean = line.lstrip("-•–— ")
            if clean and clean not in relevant:
                relevant.append(clean)
                if len(relevant) >= max_lines:
                    break
    return relevant


def extract_context_snippet(company_context, max_chars=180):
    """Return a short, usable snippet from company_context."""
    if not company_context or not company_context.strip():
        return ""
    text = company_context.strip()
    sentences = re.split(r"(?<=[.!?])\s+", text)
    snippet = ""
    for sent in sentences:
        if len(snippet) + len(sent) + 1 <= max_chars:
            snippet = (snippet + " " + sent).strip()
        else:
            break
    return snippet or text[:max_chars].rstrip()


# ── Cover letter ──────────────────────────────────────────────────────────────

def cover_letter_template(company, role, location, name, matched,
                          resume_text="", company_context=""):
    today_str = date.today().strftime("%B %d, %Y")
    loc_phrase = f" based in {location}" if location else ""

    # Opening
    opening = (
        f"I am writing to express my strong interest in the {role} position "
        f"at {company}{loc_phrase}. After reviewing the role, I am excited by "
        f"the opportunity to contribute to your team."
    )

    # Body 1: company-specific if context available
    if company_context and company_context.strip():
        snippet = extract_context_snippet(company_context, 160)
        body1 = (
            f"I was particularly drawn to {company} because of "
            f"{snippet.rstrip('.,')[:140]}. "
            f"This aligns closely with my academic background and project experience."
        )
    else:
        body1 = (
            f"I am drawn to this role because of the opportunity to apply my skills "
            f"in a team-oriented environment and contribute to meaningful work."
        )

    # Body 2: relevant resume evidence
    lines = extract_relevant_resume_lines(resume_text, matched, max_lines=3)
    if lines:
        bullets = "\n".join(f"  • {l}" for l in lines)
        body2 = (
            f"Some experiences relevant to this role include:\n{bullets}\n\n"
            + (
                f"These have strengthened my skills in "
                f"{', '.join(matched[:4])}{',' if len(matched) > 4 else ''}"
                f"{' and more' if len(matched) > 4 else ''}."
                if matched else ""
            )
        ).strip()
    elif matched:
        body2 = (
            f"My background includes practical experience in "
            f"{', '.join(matched[:5])}"
            f"{', among other areas' if len(matched) > 5 else ''}, "
            f"which align with the key requirements of this role."
        )
    else:
        body2 = (
            "My coursework, projects, and prior experience have prepared me to contribute "
            "effectively to a fast-paced, team-oriented environment."
        )

    # Closing
    closing = (
        f"Thank you for considering my application. I would welcome the chance to discuss "
        f"how my background aligns with {company}'s needs. "
        f"I am happy to provide additional materials upon request."
    )

    return (
        f"{today_str}\n\n"
        f"Dear {company} Hiring Team,\n\n"
        f"{opening}\n\n"
        f"{body1}\n\n"
        f"{body2}\n\n"
        f"{closing}\n\n"
        f"Sincerely,\n{name}\n"
    )


# ── LinkedIn message ──────────────────────────────────────────────────────────

def linkedin_template(company, role, name, matched, company_context=""):
    skill_phrase = f" with a background in {matched[0]}" if matched else ""
    if company_context and company_context.strip():
        opener = f"I've been following {company}'s work and"
    else:
        opener = f"I came across the {role} opening at {company} and"

    return (
        f"Hi [First Name],\n\n"
        f"{opener} wanted to reach out. I'm a student"
        f"{skill_phrase} interested in the {role} role. "
        f"I'd love to hear about your experience on the team — "
        f"would you be open to a quick 15-minute chat?\n\n"
        f"Thank you for your time,\n{name}\n"
    )


# ── Recruiter email ───────────────────────────────────────────────────────────

def recruiter_email_template(company, role, name, matched, company_context=""):
    if matched:
        qual_list = matched[:3]
        quals = (
            f"My background includes {', '.join(qual_list)}"
            f"{', and related skills' if len(matched) > 3 else ''}, "
            f"which align with the qualifications listed for this role."
        )
    else:
        quals = (
            "My coursework and projects have prepared me for "
            "the responsibilities outlined in the job description."
        )

    return (
        f"Subject: {role} Application — {name}\n\n"
        f"Dear {company} Recruiting Team,\n\n"
        f"I recently applied for the {role} position at {company} and wanted to briefly introduce myself. "
        f"{quals}\n\n"
        f"I have attached my resume and would welcome the chance to discuss my application. "
        f"Please let me know if you need any additional information.\n\n"
        f"Thank you for your time.\n\n"
        f"Best regards,\n{name}\n"
    )


# ── Resume suggestions ────────────────────────────────────────────────────────

RESUME_DISCLAIMER = (
    "⚠️  IMPORTANT: Only use these suggestions if they accurately reflect "
    "your real experience. Do not fabricate skills, tools, metrics, or responsibilities."
)


def resume_suggestions_template(company, role, missing, matched, resume_text=""):
    sep = "=" * 65
    header = f"Resume Improvement Suggestions\n{role} at {company}\n{sep}\n"

    disclaimer = f"\n{RESUME_DISCLAIMER}\n"

    # Strengths
    if matched:
        strengths = (
            f"\nCURRENT STRENGTHS FOR THIS ROLE\n"
            f"Your resume already shows: {', '.join(matched)}.\n"
            f"Ensure these appear in your Skills section and are each backed by a concrete bullet.\n"
        )
    else:
        strengths = (
            "\nCURRENT STRENGTHS FOR THIS ROLE\n"
            "No matching skills found in your resume for this job description.\n"
            "Consider roles that better match your current profile, or build toward this role.\n"
        )

    # Gaps
    if not missing:
        gaps = (
            "\nSKILL GAPS\n"
            "All recognized skills from the job description appear in your resume. "
            "Focus on ensuring each has a concrete supporting bullet.\n"
        )
        bullets = ""
    else:
        gaps = (
            f"\nSKILL GAPS  ({len(missing)} skills from JD not found in resume)\n"
            + "\n".join(f"  • {s}" for s in missing) + "\n"
        )
        suggestions = []
        for skill in missing[:5]:
            in_resume = any(skill.lower() in l.lower() for l in (resume_text or "").splitlines())
            if in_resume:
                suggestions.append(
                    f"  • {skill}: Appears to be in your resume but not clearly labeled. "
                    f"Add it explicitly to your Skills section."
                )
            else:
                suggestions.append(
                    f"  • {skill}: Not found in resume. If you have real experience, add a bullet. "
                    f"If not, a short course or project could address this gap."
                )
        bullets = "\nSUGGESTED IMPROVEMENTS\n" + "\n".join(suggestions) + "\n"

    tips = (
        "\nGENERAL TIPS\n"
        "  1. Quantify impact where possible (e.g., 'reduced runtime by 30%').\n"
        "  2. Use strong action verbs: Built, Analyzed, Designed, Led, Improved.\n"
        "  3. Tailor your resume objective/summary to each role if you use one.\n"
        "  4. Re-run Analyze Match after updating your resume to check your new score.\n"
    )

    return header + disclaimer + strengths + gaps + bullets + tips


# ── Follow-up and thank-you messages ─────────────────────────────────────────

def followup_email_template(company, role, name, days_since_applied=None):
    days_phrase = (
        f", submitted approximately {days_since_applied} days ago,"
        if days_since_applied else ""
    )
    return (
        f"Subject: Following Up — {role} Application\n\n"
        f"Dear {company} Recruiting Team,\n\n"
        f"I wanted to follow up on my application for the {role} position{days_phrase}. "
        f"I remain very interested in this opportunity and would welcome the chance "
        f"to discuss my qualifications further.\n\n"
        f"Please let me know if you need any additional information from me.\n\n"
        f"Thank you for your time.\n\n"
        f"Best regards,\n{name}\n"
    )


def linkedin_followup_template(company, role, name):
    return (
        f"Hi [First Name],\n\n"
        f"I wanted to follow up on my application for the {role} role at {company}. "
        f"I'm still very interested and would love any update on the timeline. "
        f"Please let me know if there is anything else I can provide.\n\n"
        f"Thank you,\n{name}\n"
    )


def thankyou_template(company, role, contact_name, name):
    contact = contact_name.strip() if contact_name and contact_name.strip() else "the team"
    return (
        f"Subject: Thank You — {role} Interview at {company}\n\n"
        f"Dear {contact},\n\n"
        f"Thank you for taking the time to speak with me about the {role} position at {company}. "
        f"I genuinely enjoyed learning about [specific topic from your conversation] "
        f"and I am even more enthusiastic about this opportunity after our chat.\n\n"
        f"[Add one specific detail from your call here — this makes the note memorable.]\n\n"
        f"I look forward to hearing about the next steps and am happy to provide "
        f"any additional information.\n\n"
        f"Best regards,\n{name}\n"
    )


# ── DOCX writer ───────────────────────────────────────────────────────────────

def write_docx(filepath, title, content):
    """Write content to a DOCX file. Returns True on success, False if unavailable."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title heading
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            # Separator lines — skip
            if re.match(r"^[=\-]{5,}$", block):
                continue
            # Disclaimer / warning lines
            if block.startswith("⚠️") or block.startswith("IMPORTANT:"):
                para = doc.add_paragraph()
                run = para.add_run(block)
                run.bold = True
                try:
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                except Exception:
                    pass
            elif "\n" in block:
                for line in block.splitlines():
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)
            else:
                doc.add_paragraph(block)

        doc.save(filepath)
        return True
    except ImportError:
        return False
    except Exception:
        return False


def write_txt(filepath, content):
    with open(filepath, "w") as f:
        f.write(content)


# ── Report generators (for Export section) ───────────────────────────────────

def generate_report_docx(title, sections, output_path):
    """Write a multi-section DOCX report. sections = [(heading, body_text), ...]"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(title, 0)
        for heading, body in sections:
            if heading:
                doc.add_heading(heading, 2)
            if body:
                for line in body.splitlines():
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)
        doc.save(output_path)
        return True
    except ImportError:
        return False
    except Exception:
        return False


# ── Main file generator ───────────────────────────────────────────────────────

def generate_files(job, matched, missing, name, output_dir, resume_text=""):
    """Generate DOCX (primary) and TXT (fallback) for all four material types.

    Returns list of written file paths.
    """
    company  = job.get("company") or "Unknown_Company"
    role     = job.get("role") or "Unknown_Role"
    location = job.get("location") or ""
    company_context = job.get("company_context") or ""

    os.makedirs(output_dir, exist_ok=True)
    base = safe_filename(f"{company}_{role}")
    written = []

    templates = {
        "Cover_Letter": cover_letter_template(
            company, role, location, name, matched, resume_text, company_context
        ),
        "LinkedIn_Message": linkedin_template(
            company, role, name, matched, company_context
        ),
        "Recruiter_Email": recruiter_email_template(
            company, role, name, matched, company_context
        ),
        "Resume_Suggestions": resume_suggestions_template(
            company, role, missing, matched, resume_text
        ),
    }

    for doc_type, content in templates.items():
        docx_path = os.path.join(output_dir, f"{base}_{doc_type}.docx")
        txt_path  = os.path.join(output_dir, f"{base}_{doc_type}.txt")
        title_str = f"{doc_type.replace('_', ' ')} — {company} — {role}"

        ok = write_docx(docx_path, title_str, content)
        write_txt(txt_path, content)    # always keep TXT as fallback

        written.append(docx_path if ok else txt_path)

    return written
